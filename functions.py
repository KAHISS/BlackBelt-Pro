# interface libs =================================
from largeVariables import *
from tkinter import messagebox
from customtkinter import *
from tkinter import filedialog
from tkinter import *
from tkinter import TclError
from tkinter.colorchooser import askcolor
from PIL import Image
import re
from docx import Document
from docx.shared import Pt

# functions libs ==================================
import os
import requests
from datetime import datetime
from reportlab.lib.pagesizes import A4
from reportlab.platypus import SimpleDocTemplate, Table, TableStyle
import shutil


class GeneralFunctions:

    def backup_dataBaes(self):
        # pick directory if origin =========================
        origin = './resources'
        destiny = os.path.join(os.path.expanduser("~"), "KonectSys/backup/resources")
        # coping ================================
        if os.path.exists(destiny):
            shutil.rmtree(destiny)
            shutil.copytree(origin, destiny)
        else:
            shutil.copytree(origin, destiny)
        self.message_window(1, 'Concluído', messagein=f'Backup feito com sucesso')

    @staticmethod
    def backup_dataBaes_discret():
        # pick directory if origin =========================
        origin = './resources'
        destiny = os.path.join(os.path.expanduser("~"), "KonectSys/backup/resources")
        # coping ================================
        if os.path.exists(destiny):
            shutil.rmtree(destiny)
            shutil.copytree(origin, destiny)
        else:
            shutil.copytree(origin, destiny)

    def loading_database(self):
        # pick directory if origin =========================
        origin = os.path.join(os.path.expanduser("~"), "KonectSys/backup/resources")
        destiny = './resources'
        # coping ================================
        if os.path.exists(origin):
            shutil.rmtree(destiny)
            shutil.copytree(origin, destiny)
        self.message_window(1, 'Concluído', messagein=f'Carregamento do backup feito com sucesso')

    @staticmethod
    def beltColor(color):
        colors = {
            'BRANCA': 'branca',
            'CINZA': 'cinza',
            'AMARELA': 'amarela',
            'LARANJA': 'laranja',
            'VERDE': 'verde',
            'AZUL': 'azul',
            'ROXA': 'roxa',
            'MARROM': 'marrom',
            'PRETA': 'preta'
        }
        return colors[color]

    def insert_treeview_informations(self, treeview, infos, line_color):
        for info in infos:
            if self.lineTreeviewColor[line_color] % 2 == 0:
                treeview.insert('', 'end', values=info, tags='oddrow')
            else:
                treeview.insert('', 'end', values=info, tags='evenrow')
            self.lineTreeviewColor[line_color] += 1

    def insert_treeview_informations_oneLine(self, treeview, infos, line_color):
        if treeview == self.treeviewSchedule:
            for info in infos:
                treeview.insert(
                    '', 'end',
                    values=info,
                    tags=self.beltColor(self.dataBases['informations'].searchDatabase(f'SELECT faixa FROM Alunos WHERE nome = "{info[1]}"')[0][0]))
        else:
            for info in infos:
                treeview.insert('', 'end', values=info, tags=self.beltColor(info[5]))

    @staticmethod
    def pick_informations_treeview(treeview):
        selection = treeview.get_children()
        information = []
        for i in selection:
            information.append(treeview.item(i, 'values'))
        return information

    @staticmethod
    def selection_treeview(treeview):
        selection = treeview.selection()
        information = []
        for i in selection:
            information.append(treeview.item(i, 'values'))
        return information

    @staticmethod
    def request_adrees(zip_code, informations):
        # treating the cep =============================
        treatedZipCode = zip_code.replace('-', '').replace(' ', '').replace('.', '')
        # validating if the zip code have eight numbers =======================
        if len(treatedZipCode) == 8:
            try:
                # making request in the site ========================
                request = requests.get(f'http://viacep.com.br/ws/{treatedZipCode}/json/').json()
            except Exception:
                # error in request ===================
                pass
            else:
                if len(request) > 1:
                    # deleting informations in the entrys =====================
                    for information in informations:
                        information.delete(0, END)
                    # insert informations in the entrys ========================
                    informations[0].insert(0, request['bairro'])
                    informations[1].insert(0, request['localidade'])
                    informations[2].insert(0, request['uf'])

    @staticmethod
    def validation(infos, type_validation, index=None):
        if type_validation == 1:
            for info in infos[0]:
                if info == '' or info == 'R$,00':
                    return False
        elif type_validation == 2:
            if not infos.replace('.', '', 1).isdigit():
                return False
        elif type_validation == 3:
            for c in infos:
                if c.isalpha() or c == '.':
                    return False
        elif type_validation == 4:
            for c in infos:
                if c.isalpha() or c == '.':
                    return False
        elif type_validation == 5:
            for info in infos:
                if info == '' or info == 'R$,00' or info == ':00':
                    return False
        elif type_validation == 6:
            for c in infos:
                if c.isalpha() or c == '.':
                    return False
            if len(infos) != 12:
                return False
        elif type_validation == 7:
            for info in infos:
                for c in info:
                    if c.isalpha():
                        return False
        elif type_validation == 8:
            for info in infos:
                for c in info[2:]:
                    if c.isalpha() or c == '.':
                        return False
        return True

    @staticmethod
    def treating_numbers(info=None, type_treating=1, values=None, entry2=None, ide=4):
        if type_treating == 1:
            if ',' in info:
                value = info.replace('R$', '').strip().split(',')
                if value[1] == '':
                    return 'R$' + ','.join(value) + '00'
                else:
                    return 'R$' + ','.join(value)
            else:
                value = info.replace('R$', '').strip()
                return 'R$' + value + ',00'

        elif type_treating == 2:
            sum_value = 0
            for value in values:
                number = float(value[ide].replace('R$', '').replace(',', '.'))
                sum_value += number
            return 'R$' + f'{sum_value:.2f}'.replace('.', ',')

        elif type_treating == 3:
            if ':' in info:
                hour = info.split(':')
                if hour[1] == '':
                    hour[1] = '00'
                if hour[0] == '':
                    hour[0] = '00'
                if len(hour[0]) == 1:
                    hour[0] = '0' + hour[0]
                if hour[0] and hour[1] == '':
                    hour[0], hour[1] = '00'
                return ':'.join(hour)
            else:
                if len(info) == 1:
                    info = '0' + info
                return info + ':00'

        elif type_treating == 4:
            sum_value = 0
            for value in values:
                number = float(value.replace('R$', '').replace(',', '.'))
                sum_value += number
            return 'R$' + f'{sum_value:.2f}'.replace('.', ',')

        elif type_treating == 5:
            subtraction_value = float(values[0].replace('R$', '').replace(',', '.'))
            for value in values[1:]:
                number = float(value.replace('R$', '').replace(',', '.'))
                subtraction_value -= number
            return 'R$' + f'{subtraction_value:.2f}'.replace('.', ',')

        elif type_treating == 6:
            sum_value = 0
            for value in values:
                number = float(value.replace('R$', '').replace(',', '.'))
                sum_value += number
            return int(sum_value)

        elif type_treating == 7:
            number = float(info.replace('R$', '').replace(',', '.'))
            return number

        elif type_treating == 8:
            phone = info.get().replace(' ', '').replace('-', '').replace('(', '').replace(')', '')
            phoneFormated = ''
            if len(info.get()) == 9:
                phoneFormated = f'(77) {phone[0:5]}-{phone[5:]}'
            if len(info.get()) >= 11:
                phoneFormated = f'({phone[0:2]}) {phone[2:7]}-{phone[7:]}'
            info.delete(0, END)
            info.insert(0, phoneFormated)

        elif type_treating == 9:
            phone = info.get().replace(' ', '').replace('-', '').replace('(', '').replace(')', '')
            phoneFormated = ''
            if len(info.get()) == 9:
                phoneFormated = f'(77) {phone[0:5]}-{phone[5:]}'
            if len(info.get()) >= 11:
                phoneFormated = f'({phone[0:2]}) {phone[2:7]}-{phone[7:]}'
            info.delete(0, END)
            info.insert(0, phoneFormated)

        elif type_treating == 10:
            cpf = info.get().replace(' ', '').replace('-', '').replace('.', '')
            cpfFormated = ''
            if len(info.get()) == 11:
                cpfFormated = f'{cpf[0:3]}.{cpf[3:6]}.{cpf[6:9]}-{cpf[9:]}'
            else:
                cpfFormated = info.get()
            info.delete(0, END)
            info.insert(0, cpfFormated)
        elif type_treating == 11:
            rg = info.get().replace(' ', '').replace('-', '').replace('.', '')
            rgFormated = ''
            if len(info.get()) == 9:
                rgFormated = f'{rg[0:2]}.{rg[2:5]}.{rg[5:8]}-{rg[8]}'
            else:
                rgFormated = info.get()
            info.delete(0, END)
            info.insert(0, rgFormated)

    def delete_informations_treeview(self, treeview, line_color):
        for linhas in treeview.get_children():
            treeview.delete(linhas)
        self.lineTreeviewColor[line_color] = 0

    @staticmethod
    def message_window(typem=1, titlein='', messagein=''):
        if typem == 1:
            messagebox.showinfo(title=titlein, message=messagein)
        elif typem == 2:
            messagebox.showwarning(title=titlein, message=messagein)
        elif typem == 3:
            messagebox.showerror(title=titlein, message=messagein)
        elif typem == 4:
            question = messagebox.askyesno(title=titlein, message=messagein)
            return question

    def insert_informations_entrys(self, entrys, treeview=None, insert=True, type_insert='normal', table='', photo=None, size=(170, 200), data_base='informations'):
        match type_insert:
            case 'normal':
                # deleting informations of entrys ===============================
                for entry in entrys:
                    if isinstance(entry, CTkComboBox):
                        entry.set('')
                    else:
                        entry.delete(0, END)
                # cheking if there is information in the treeview ================================
                if insert:
                    if treeview.selection():
                        for index, information in enumerate(self.selection_treeview(treeview)[0][1:]):
                            if isinstance(entrys[index], CTkComboBox):
                                entrys[index].set(information)
                            else:
                                entrys[index].insert(0, information)
                                if entrys[index] == self.passwordEntry:
                                    entrys[index].delete(0, END)
            case 'advanced':
                # deleting informations of entrys ===============================
                for entry in entrys:
                    if isinstance(entry, CTkComboBox) or isinstance(entry, StringVar):
                        entry.set('')
                    elif isinstance(entry, CTkLabel):
                        self.pick_picture(entry, photo, 'unknow')
                    elif isinstance(entry, CTkTextbox):
                        entry.delete('0.0', END)
                    else:
                        entry.delete(0, END)
                if insert:
                    # cheking if there is information in the treeview ================================
                    if treeview.selection():
                        informations = self.selection_treeview(treeview)[0]
                        # case informations is of cash day
                        if 'R$' in informations[2]:
                            informations = informations[0:6] + informations[10:13] + ('', '', '')
                            print(informations)
                        for index, information in enumerate(informations[1:]):
                            if isinstance(entrys[index], CTkComboBox) or isinstance(entrys[index], StringVar):
                                entrys[index].set(information.capitalize() if isinstance(entrys[index], StringVar) else information)
                            elif isinstance(entrys[index], CTkLabel):
                                self.pick_picture(entrys[index], photo, 'toggle', self.dataBases[data_base].searchDatabase(f'SELECT foto FROM {table} WHERE ID = {informations[0]}')[0][0], size=size)
                            elif isinstance(entrys[index], CTkTextbox):
                                entrys[index].insert('0.0', self.dataBases[data_base].searchDatabase(f'SELECT observação FROM {table} WHERE ID = {informations[0]}')[0][0])
                            else:
                                entrys[index].insert(0, information)

    def delete_information(self, treeview, type_information, table):
        # pick up selection for delete information =============================
        if ask := self.message_window(4, 'Comfimação', 'Você tem certeza de que deseja deletar o(s) item(s) selecionado(s)'):
            for selection in self.selection_treeview(treeview):
                self.dataBases[type_information].crud(deleteInformation.format(table, selection[0]))
            if len(self.dataBases[type_information].searchDatabase(searchAll.format(table))) == 0:
                self.dataBases[type_information].crud(f"DELETE FROM sqlite_sequence WHERE name='{table}'")
        return ask

    @staticmethod
    def create_pdf(treeview, elements):
        saveDirectory = filedialog.asksaveasfilename(defaultextension="*.pdf ", filetypes=[('Arquivos de pdf', "*.pdf")])
        if saveDirectory != '':
            # creating documents ============================================
            document = SimpleDocTemplate(saveDirectory, pagesize=A4)
            # insert elements in documents ==================================
            content = []
            for element in elements:
                content.append(element)
            document.build(elements)
            return saveDirectory

    def create_record(self, informations):
        saveDirectory = filedialog.askdirectory()
        if saveDirectory != '':
            # pick document of origin
            tokenOfStudent = Document('documents/ficha_do_aluno.docx')
            if len(informations) > 0:
                for information in informations:
                    # insert elements in references ==================================
                    index = 1
                    for key in references:
                        references[key] = information[index]
                        index += 1
                    # replacement itens =========================
                    for paragraph in tokenOfStudent.paragraphs:
                        paragraph.style.font.name = 'Arial'
                        paragraph.style.font.size = Pt(16)
                        for key, value in references.items():
                            if key in paragraph.text:
                                paragraph.text = paragraph.text.replace(key, value)
                    # saving ========================
                    tokenOfStudent.save(saveDirectory + f'/{information[1]}.docx')
                self.message_window(1, 'Concluído', messagein=f'Os arquivos foram salvos em "{saveDirectory}"')
            else:
                # insert elements in references ==================================
                for key in references:
                    if key == 'pspsps' or key == 'qqq':
                        references[key] = '_' * 141
                    else:
                        references[key] = '_' * 47
                # replacement itens =========================
                for paragraph in tokenOfStudent.paragraphs:
                    paragraph.style.font.name = 'Arial'
                    paragraph.style.font.size = Pt(16)
                    for key, value in references.items():
                        if key in paragraph.text:
                            paragraph.text = paragraph.text.replace(key, value)
                # saving ========================
                tokenOfStudent.save(saveDirectory + f'/ficha_do_aluno.docx')
                self.message_window(1, 'Concluído', messagein=f'Os arquivos foram salvos em "{saveDirectory}"')

    @staticmethod
    def image(file, size):
        try:
            if 'icon_no_picture.png' in file or 'icon_barCode.png' in file or 'icon_product.png' in file:
                size = (76, 76)
            img = CTkImage(light_image=Image.open(file), dark_image=Image.open(file), size=size)
        except FileNotFoundError:
            img = CTkImage(light_image=Image.open('assets/corrupted.png'), dark_image=Image.open('assets/corrupted.png'), size=(76, 76))
        return [img, file]

    def active_monthly_payments(self, treeview):
        self.search_student(treeview, type_search='all', save_seacrh=False)
        # excluing selections ================
        for item in treeview.selection():
            treeview.selection_remove(item)

        # add selections ================
        for student in treeview.get_children():
            last_date = re.findall(date_pattern, treeview.set(student, "Observações"))
            if last_date:
                if datetime.today().strftime('%m/%Y') != last_date[0]:
                    treeview.selection_add(student)
            else:
                treeview.selection_add(student)

    def pick_picture(self, label, photo, type_photo='new', directory='', size=None):
        # pick directory of photo ======================================
        match type_photo:
            case 'new':
                fileName = filedialog.askopenfilename()
                if fileName:
                    self.photosAndIcons[photo] = self.image(fileName, (170, 200))
                    label.configure(image=self.photosAndIcons[photo][0])
            case 'toggle':
                self.photosAndIcons[photo] = self.image(directory, size)
                label.configure(image=self.photosAndIcons[photo][0])
            case 'unknow':
                if photo in ['employee', 'costumer']:
                    self.photosAndIcons[photo] = self.image(f'assets/icon_no_picture.png', (76, 76))
                elif photo == 'barCode':
                    self.photosAndIcons[photo] = self.image(f'assets/icon_barCode.png', (76, 76))
                elif photo in ['productUse', 'productSale', 'productUseUnusable', 'productSaleSold']:
                    self.photosAndIcons[photo] = self.image(f'assets/icon_product.png', (76, 76))
                label.configure(image=self.photosAndIcons[photo][0])
            case 'logo':
                fileName = filedialog.askopenfilename()
                if fileName:
                    if '.png' in fileName.lower():
                        label.configure(image=self.image(fileName, (500, 500))[0])
                        self.dataBases['config'].crud(f'UPDATE Logo SET arquivo="{fileName}"')
                    else:
                        self.message_window(3, 'Formato incorreto', 'Use uma imagem em formato PNG')

    @staticmethod
    def searching_list(first, quantity, column, insert=False, index=0, information=''):
        # create list of informations ===============================
        listSearch = [first]
        listSearch.extend([''] * quantity)
        listSearch.append(column)
        if insert:
            listSearch.insert(index, information)
        return listSearch

    def encode_for_searching(self, information):
        if information == '':
            return information
        else:
            return self.criptography.encode(information)

    def decode_informations_database(self, informations):
        # decoding informations ================================
        informationsDecode = []
        for information in informations:
            informationsDecode.append(
                (
                    information[0], self.criptography.decode(information[1][2:-1]), self.criptography.decode(information[2][2:-1]), self.criptography.decode(information[3][2:-1]),
                    self.criptography.decode(information[4][2:-1]), self.criptography.decode(information[5][2:-1]), self.criptography.decode(information[6][2:-1]), self.criptography.decode(information[7][2:-1]),
                    self.criptography.decode(information[8][2:-1]), self.criptography.decode(information[9][2:-1]), self.criptography.decode(information[10][2:-1]), self.criptography.decode(information[11][2:-1]),
                    information[11].upper(), information[13].upper()
                )
            )
        return informationsDecode


class FunctionsOfSchedule(GeneralFunctions):

    def register_scheduling(self, informations, treeview, entrys=None):
        # informations of treeview ====================
        if self.validation(informations[0:4], 5) and self.validation(informations[1], 4) and self.validation(informations[3], 3):
            if self.message_window(4, 'Comfimação', 'Você tem certeza de que deseja finalizar o(s) agendamento(s)'):
                self.dataBases['schedule'].crud(
                    registerScheduling.format(
                        informations[0].title(), self.treating_numbers(informations[1], 1), informations[2].upper(),
                        datetime.today().strftime('%d/%m/%Y') if informations[3] == '' else informations[3].upper(), datetime.today().strftime('%m/%Y')
                    ))
                self.search_schedule(treeview, informations, 'last', save_seacrh=False)

                # anoting payment ================
                self.dataBases['informations'].crud(f'UPDATE Alunos SET observação = "Mensalidade paga:\n{datetime.today().strftime("%m/%Y")}" WHERE nome = "{informations[0].title()}"')

                self.message_window(1, 'Concluído', messagein=f'Agendamento(s) feito com sucesso')
        else:
            self.message_window(3, 'Erro', 'Verifique se os campos estão preenchidos ou corretos')

    def search_schedule(self, treeview=None, informations=None, type_search='new', save_seacrh=True, insert=True):
        # save last search ============================================
        if save_seacrh:
            self.lastSearch['schedule'] = searchSchedule.format(
                'ID' if informations[0].isnumeric() else 'aluno', informations[0], informations[1], informations[2], informations[3], informations[4],
                informations[5].replace(' ', '_').lower()
            )

        # pick up informations =========================================
        informationsDataBase = []
        match type_search:
            case 'new':
                informationsDataBase = self.dataBases['schedule'].searchDatabase(
                    searchSchedule.format(
                        'ID' if informations[0].isnumeric() else 'aluno', informations[0], informations[1], informations[2], informations[3], informations[4],
                        informations[5].replace(' ', '_').lower()
                    )
                )
            case 'last':
                informationsDataBase = self.dataBases['schedule'].searchDatabase(self.lastSearch['schedule'])
            case 'all':
                informationsDataBase = self.dataBases['schedule'].searchDatabase(searchAll.format('Agenda'))

        if insert:
            # deleting and inserting informations in treeview ===============================
            self.delete_informations_treeview(treeview, 'schedule')
            self.insert_treeview_informations_oneLine(treeview, informationsDataBase, 'schedule')
        else:
            return informationsDataBase

    def update_schedule(self, treeview, informations, entrys):
        # update informations =========================================
        if self.validation(informations, 5) and self.validation(informations[1], 4) and self.validation(informations[3], 3):
            if treeview.selection():
                informationsDataBase = self.dataBases['schedule'].crud(
                    updateSchedule.format(
                        informations[0].title(), self.treating_numbers(informations[1], 1), informations[2].upper(), informations[3].upper(), informations[4].upper(),
                        self.selection_treeview(treeview)[0][0]
                    )
                )
                # delete informations of treeview ==============================
                self.delete_informations_treeview(treeview, 'schedule')

                # insert informations in treeview ===============================
                self.search_schedule(treeview, informations, 'last', save_seacrh=False)

                # show message of concluded
                self.message_window(1, 'Concluído', messagein=f'agendamento(s) atualizado(s) com sucesso')
            else:
                self.message_window(3, 'Sem seleção', 'Selecione algum item na lista para atualizar')
        else:
            self.message_window(3, 'Erro', 'Verifique se os campos estão preenchidos ou corretos')

    def delete_schedule(self, treeview):
        if treeview.selection():
            # deleting inforations =======================================
            delete = self.delete_information(treeview, 'schedule', 'Agenda')

            if delete:
                # delete informations of treeview ==============================
                self.delete_informations_treeview(treeview, 'schedule')

                # insert informations in treeview ===============================
                self.search_schedule(treeview, type_search='last', save_seacrh=False)

                # shoe message of concluded
                self.message_window(1, 'Concluído', messagein=f'agendamento(s) deletado(s) com sucesso')
        else:
            self.message_window(3, 'Sem seleção', 'Selecione algum item na lista para deletar')

    def create_pdf_schedule(self, treeview):
        informationsTreeview = self.pick_informations_treeview(treeview)
        if informationsTreeview:
            # Collecting informations for messege===============================
            amountClients = len(informationsTreeview)
            sumValue = self.treating_numbers(type_treating=2, values=informationsTreeview, ide=2)
            methoPay = {
                'card': [row for row in informationsTreeview if row[3] in ['CARTÃO', 'CARTAO']],
                'money': [row for row in informationsTreeview if row[3] in ['DINHEIRO']],
                'transfer': [row for row in informationsTreeview if row[3] in ['TRANSFERÊNCIA', 'TRANSFERENCIA']],
                'note': [row for row in informationsTreeview if row[3] in ['NOTA', 'FIADO', 'NOTINHA']],
                'notPay': [row for row in informationsTreeview if row[3] in ['NÃO FOI PAGO', 'SEM PAGAMENTO', 'NÃO PAGO', '']],
            }

            sumMetohdPay = [
                amountClients,
                self.treating_numbers(type_treating=2, values=methoPay["card"], ide=2),
                self.treating_numbers(type_treating=2, values=methoPay["money"], ide=2),
                self.treating_numbers(type_treating=2, values=methoPay["transfer"], ide=2),
                self.treating_numbers(type_treating=2, values=methoPay["note"], ide=2),
                self.treating_numbers(type_treating=2, values=methoPay["notPay"], ide=2),
                sumValue
            ]

            # informations ============================================
            for information in informationsTreeview:
                tableWithInformationsScheduleTreeview.append(information)
            tableWithInformationsComplementarySchedule.append(sumMetohdPay)

            # create tables ===========================================
            table1 = Table(tableWithInformationsScheduleTreeview)
            table1.setStyle(TableStyle(styleTableInformationsTreeview))
            table2 = Table(tableWithInformationsComplementarySchedule)
            table2.setStyle(TableStyle(styleTableInformationsComplementary))

            # reseting tables =========================================
            del tableWithInformationsScheduleTreeview[2:]
            del tableWithInformationsComplementarySchedule[2:]

            # creating pdf ============================================
            saveDirectory = self.create_pdf(treeview, [table1, table2])
            if saveDirectory is not None:
                self.message_window(1, 'Concluído', messagein=f'O arquivo foi salvo em "{saveDirectory}"')
        else:
            self.message_window(2, 'Sem registro', 'A tabela esta vazia')

    def message_informations_schedule(self, treeview):
        if informationsTreeview := self.pick_informations_treeview(treeview):
            # Collecting informations for messege===============================
            amountClients = len(informationsTreeview)
            sumValues = self.treating_numbers(type_treating=2, values=informationsTreeview, ide=2)
            methoPay = {
                'card': [row for row in informationsTreeview if row[3] in ['CARTÃO', 'CARTAO']],
                'money': [row for row in informationsTreeview if row[3] in ['DINHEIRO']],
                'transfer': [row for row in informationsTreeview if row[3] in ['TRANSFERÊNCIA', 'TRANSFERENCIA']],
                'note': [row for row in informationsTreeview if row[3] in ['NOTA', 'FIADO', 'NOTINHA']],
                'notPay': [row for row in informationsTreeview if row[3] in ['NÃO FOI PAGO', 'SEM PAGAMENTO', 'NÃO PAGO', '']],
            }

            # shoe menssege ================================================
            self.message_window(
                1,
                'Informações sobre a tabela',
                f'Total de Alunos = {amountClients}\n'
                f'Total em cartão = {self.treating_numbers(type_treating=2, values=methoPay["card"], ide=2)}\n'
                f'Total em dinheiro = {self.treating_numbers(type_treating=2, values=methoPay["money"], ide=2)}\n'
                f'Total em tranferència = {self.treating_numbers(type_treating=2, values=methoPay["transfer"], ide=2)}\n'
                f'Total em nota = {self.treating_numbers(type_treating=2, values=methoPay["note"], ide=2)}\n'
                f'Total não pago = {self.treating_numbers(type_treating=2, values=methoPay["notPay"], ide=2)}\n'
                f'Total recebido = {sumValues}'
            )
        else:
            self.message_window(2, 'Sem registro', 'A tabela esta vazia')


class FunctionsOfStudentInformations(GeneralFunctions):

    def register_student(self, informations, treeview):
        if self.validation(informations[0:19], 5) and self.validation([informations[1], informations[2], informations[3], informations[10], informations[11], informations[12], informations[14]], 7):
            if self.message_window(4, 'Comfimação', f'Finalisar o cadastro de {informations[0].title()}?'):
                # informations of treeview ====================
                self.dataBases['informations'].crud(
                    registerStudent.format(
                        informations[0].title(), informations[1].upper(), informations[2].upper(), informations[3].upper(), informations[4].upper(), informations[5].upper(),
                        informations[6].upper(), informations[7].upper() if informations[7] != '' and informations[6] != 'Sim' else 'NENHUM PROBLEMA',
                        informations[8].upper(), informations[9].upper() if informations[9] != '' and informations[8] != 'Não' else 'NENHUMA LIMITAÇÂO', informations[10].upper(), informations[11].upper(),
                        informations[12].upper(), informations[13].upper(), informations[14].upper(), informations[15].upper(), informations[16].upper(), informations[17].upper(), informations[18],
                        f'Mensalidade em aberto'
                    )
                )
                # deleting and inserting informations in treeview ===============================
                self.search_student(treeview, informations, 'all', save_seacrh=False)

                # refresh =======================================================
                self.refresh_combobox_student()
        else:
            self.message_window(3, 'Erro', 'Verifique se os campos estão preenchidos ou corretos')

    def search_student(self, treeview=None, informations=None, type_search='new', save_seacrh=True, insert=True):
        # save last search ============================================
        if save_seacrh:
            self.lastSearch['student'] = searchStudent.format(
                'ID' if informations[0].isnumeric() else 'nome',
                informations[0].upper(), informations[1].upper(), informations[2].upper(), informations[3].upper(), informations[4].upper(), informations[5].upper(),
                informations[6].upper(), informations[7].upper(), informations[8].upper(), informations[9].upper(), informations[10].upper(), informations[11].upper(), informations[12].upper(),
                informations[13].upper(), informations[14].upper(), informations[15].upper(), informations[16].upper(), informations[17].upper(), informations[20].replace(' ', '_')
            )
        # pick up informations =========================================
        informationsDatabase = []
        match type_search:
            case 'new':
                informationsDatabase = self.dataBases['informations'].searchDatabase(
                    searchStudent.format(
                        'ID' if informations[0].isnumeric() else 'nome',
                        informations[0].upper(), informations[1].upper(), informations[2].upper(), informations[3].upper(), informations[4].upper(), informations[5].upper(),
                        informations[6].upper(), informations[7].upper(), informations[8].upper(), informations[9].upper(), informations[10].upper(), informations[11].upper(), informations[12].upper(),
                        informations[13].upper(), informations[14].upper(), informations[15].upper(), informations[16].upper(), informations[17].upper(), informations[20].replace(' ', '_')
                    )
                )
            case 'last':
                informationsDatabase = self.dataBases['informations'].searchDatabase(self.lastSearch['student'])
            case 'all':
                informationsDatabase = self.dataBases['informations'].searchDatabase(searchAll.format('Alunos'))

        if insert:
            # deleting and inserting informations in treeview ===============================
            self.delete_informations_treeview(treeview, 'student')
            self.insert_treeview_informations_oneLine(treeview, informationsDatabase, 'student')
        else:
            return informationsDatabase

    def update_student(self, treeview, informations, entrys):
        # update informations =========================================
        if treeview.selection():
            if self.validation(informations, 5) and self.validation([informations[1], informations[2], informations[3], informations[10], informations[11], informations[12], informations[14]], 7):
                informationsDataBase = self.dataBases['informations'].crud(
                    updateStudent.format(
                        informations[0].title(), informations[1].upper(), informations[2].upper(), informations[3].upper(), informations[4].upper(), informations[5].upper(),
                        informations[6].upper(), informations[7].upper(), informations[8].upper(), informations[9].upper(), informations[10].upper(), informations[11].upper(),
                        informations[12].upper(), informations[13].upper(), informations[14].upper(), informations[15].upper(), informations[16].upper(), informations[17].upper(), informations[18], self.selection_treeview(treeview)[0][0]
                    )
                )
                # delete informations of treeview ==============================
                self.delete_informations_treeview(treeview, 'student')

                # insert informations in treeview ===============================
                self.search_student(treeview, informations, 'last', save_seacrh=False)

                # refresh =======================================================
                self.refresh_combobox_student()

                # show message of concluded
                self.message_window(1, 'Concluído', messagein=f'Cadastro de {informations[0].title()} atualizado com sucesso')
            else:
                self.message_window(3, 'Erro', 'Verifique se os campos estão preenchidos ou corretos')
        else:
            self.message_window(3, 'Sem seleção', 'Selecione algum item na lista para atualizar')

    def delete_student(self, treeview):
        if treeview.selection():
            # deleting inforations =======================================
            delete = self.delete_information(treeview, 'informations', 'Alunos')

            if delete:
                # delete informations of treeview ==============================
                self.delete_informations_treeview(treeview, 'student')

                # insert informations in treeview ===============================
                self.search_student(treeview, type_search='last', save_seacrh=False)

                # refresh =======================================================
                self.refresh_combobox_student()

                # shoe message of concluded
                self.message_window(1, 'Concluído', messagein=f'Cadastro(s) deletado(s) com sucesso')
        else:
            self.message_window(3, 'Sem seleção', 'Selecione algum item na lista para deletar')

    def create_pdf_student(self, treeview):
        if informationsTreeview := self.pick_informations_treeview(treeview):
            # Collecting informations for messege===============================
            informationsComplementary = [
                len(informationsTreeview),
                max([row[3].replace(',', '.') for row in informationsTreeview], key=float),
                min([row[3].replace(',', '.') for row in informationsTreeview], key=float),
                max([row[4].replace(',', '.') for row in informationsTreeview], key=float),
                min([row[4].replace(',', '.') for row in informationsTreeview], key=float),
            ]

            # informations ============================================
            for information in informationsTreeview:
                tableWithInformationsStudentTreeview1.append(information[0:6])
            for information in informationsTreeview:
                tableWithInformationsStudentTreeview2.append(information[6:12])
            for information in informationsTreeview:
                tableWithInformationsStudentTreeview3.append(information[12:19])
            tableWithInformationsComplementaryStudent.append(informationsComplementary)

            # create tables ===========================================
            table1 = Table(tableWithInformationsStudentTreeview1)
            table1.setStyle(TableStyle(styleTableInformationsTreeview))
            table2 = Table(tableWithInformationsStudentTreeview2)
            table2.setStyle(TableStyle(styleTableInformationsTreeview))
            table3 = Table(tableWithInformationsStudentTreeview3)
            table3.setStyle(TableStyle(styleTableInformationsTreeview))
            table4 = Table(tableWithInformationsComplementaryStudent)
            table4.setStyle(TableStyle(styleTableInformationsComplementary))

            # reseting tables =========================================
            del tableWithInformationsStudentTreeview1[2:]
            del tableWithInformationsStudentTreeview2[2:]
            del tableWithInformationsStudentTreeview3[2:]
            del tableWithInformationsComplementaryStudent[2:]

            # creating pdf ============================================
            saveDirectory = self.create_pdf(treeview, [table1, table2, table3, table4])
            if saveDirectory is not None:
                self.message_window(1, 'Concluído', messagein=f'O arquivo foi salvo em "{saveDirectory}"')
        else:
            self.message_window(2, 'Sem registro', 'A tabela esta vazia')

    def message_informations_student(self, treeview):
        if informationsTreeview := self.pick_informations_treeview(treeview):
            # Collecting informations for messege===============================
            amountClients = len(informationsTreeview)
            informationsComplementary = [
                len(informationsTreeview),
                max([row[3].replace(',', '.') for row in informationsTreeview], key=float),
                min([row[3].replace(',', '.') for row in informationsTreeview], key=float),
                max([row[4].replace(',', '.') for row in informationsTreeview], key=float),
                min([row[4].replace(',', '.') for row in informationsTreeview], key=float),
            ]

            # shoe menssege ================================================
            self.message_window(
                1,
                'Informações sobre a tabela',
                f'Total de clientes = {informationsComplementary[0]}\n'
                f'Maior = {informationsComplementary[1]}\n'
                f'Menor = {informationsComplementary[2]}\n'
                f'Mais pesado = {informationsComplementary[3]}\n'
                f'Mais leve = {informationsComplementary[4]}\n'
            )
        else:
            self.message_window(2, 'Sem registro', 'A tabela esta vazia')

    def refresh_combobox_student(self):
        # refresh list of combobox services ================================
        self.customScheduleEntry.configure(values=[name[1] for name in self.search_student(informations=self.searching_list('', 19, 'nome'), save_seacrh=False, insert=False)])


class FunctionsOfCashManagement(GeneralFunctions):

    def register_cashManagement(self, informations, treeview):
        # register usage stock ====================
        if self.validation(informations[0:7], 5) and self.validation(informations[0], 7) and self.validation(informations[1:6], 8):
            if self.message_window(4, 'Comfimação', f'Registrar os dados?'):
                self.dataBases['cash'].crud(
                    registerCashManagement.format(
                        informations[0].upper(), self.treating_numbers(informations[1], 1), self.treating_numbers(informations[2], 1), self.treating_numbers(informations[3], 1),
                        self.treating_numbers(informations[4], 1), '0', '0', '0', '0', self.treating_numbers(informations[5], 1),
                        informations[6], informations[7].upper() if informations[7] != '' else 'DIA EM ANDAMENTO', informations[8].upper(),
                        informations[9]
                    ))
                # deleting and inserting informations in treeview ===============================
                self.search_cashManagement(treeview, informations, type_search='all', save_seacrh=False,)
        else:
            # show message error =====================================================
            self.message_window(3, 'Erro', 'Verifique se os campos estão preenchidos ou corretos')

    def search_cashManagement(self, treeview=None, informations=None, type_search='new', save_seacrh=True, insert=True):
        # save last search ============================================
        if save_seacrh:
            self.lastSearch['cash'] = searchCashManagement.format(
                informations[0],
                informations[1],
                informations[2],
                informations[3],
                informations[4],
                f's_{informations[10]}'.lower() if informations[10] != '' else 's_cartão',
                informations[9],
                informations[5],
                informations[6],
                informations[7],
                informations[8],
                informations[11].replace('/', '_').lower()
            )

        # pick up informations =========================================
        informationsDataBase = []
        match type_search:
            case 'new':
                informationsDataBase = self.dataBases['cash'].searchDatabase(
                    searchCashManagement.format(
                        informations[0],
                        informations[1],
                        informations[2],
                        informations[3],
                        informations[4],
                        f's_{informations[10]}'.lower() if informations[10] != '' else 's_cartão',
                        informations[9],
                        informations[5],
                        informations[6],
                        informations[7],
                        informations[8],
                        informations[11].replace('/', '_').lower(),
                    )
                )
            case 'last':
                informationsDataBase = self.dataBases['cash'].searchDatabase(self.lastSearch['cash'])
            case 'all':
                informationsDataBase = self.dataBases['cash'].searchDatabase(searchAll.format('Gerenciamento_do_mês'))

        if insert:
            # deleting and inserting informations in treeview ===============================
            self.delete_informations_treeview(treeview, 'cash')
            self.insert_treeview_informations(treeview, informationsDataBase, 'cash')
        else:
            return informationsDataBase

    def update_cashManagement(self, treeview, informations):
        # update informations =========================================
        if treeview.selection():
            if self.validation(informations[0:7], 5) and self.validation(informations[0], 7) and self.validation(informations[1:6], 8):
                # pick value for exit =======================
                exitValue = self.treating_numbers(informations[8] if informations[8] != '' else '0', 1)
                value = self.dataBases['cash'].searchDatabase(f'SELECT {"s_" + informations[9].lower() if informations[9] != "" else "s_cartão"} FROM Gerenciamento_do_mês WHERE ID = {self.selection_treeview(treeview)[0][0]}')[0][0]
                self.dataBases['cash'].crud(
                    updateCashManagement.format(
                        informations[0].upper(),
                        self.treating_numbers(informations[1], 1),
                        self.treating_numbers(informations[2], 1),
                        self.treating_numbers(informations[3], 1),
                        self.treating_numbers(informations[4], 1),
                        "s_" + informations[9].lower() if informations[9] != "" else "s_cartão",
                        self.treating_numbers(values=[value, exitValue], type_treating=4),
                        self.treating_numbers(informations[5], 1),
                        informations[6].upper(),
                        informations[7].upper(),
                        informations[10].upper(),
                        self.selection_treeview(treeview)[0][0]
                    )
                )
                # updating value received =====================================
                totalReceived = self.treating_numbers(values=[self.treating_numbers(informations[1], 1), self.treating_numbers(informations[2], 1), self.treating_numbers(informations[3], 1), self.treating_numbers(informations[4], 1)], type_treating=4)
                discount = self.treating_numbers(values=self.dataBases['cash'].searchDatabase(f'SELECT s_cartão, s_dinheiro, s_transferência, s_nota FROM Gerenciamento_do_mês WHERE ID = {self.selection_treeview(treeview)[0][0]}')[0], type_treating=4)
                self.dataBases['cash'].crud(f'UPDATE Gerenciamento_do_mês SET t_recebido = "{self.treating_numbers(values=[totalReceived, discount], type_treating=5)}" WHERE ID = {self.selection_treeview(treeview)[0][0]}')
                # delete informations of treeview ==============================
                self.delete_informations_treeview(treeview, 'cash')

                # deleting and inserting informations in treeview ===============================
                self.search_cashManagement(
                    treeview, informations, type_search='last', save_seacrh=False
                )

                # show message of concluded
                self.message_window(1, 'Concluído', messagein=f'Mês atualizado com sucesso')
            else:
                # show message error =====================================================
                self.message_window(3, 'Erro', 'Verifique se os campos estão preenchidos ou corretos')
        else:
            self.message_window(3, 'Sem seleção', 'Selecione algum item na lista para atualizar')

    def delete_cashManagement(self, treeview):
        if treeview.selection():
            # deleting inforations =======================================
            delete = self.delete_information(treeview, 'cash', 'Gerenciamento_do_mês')

            if delete:
                # delete informations of treeview ==============================
                self.delete_informations_treeview(treeview, 'cash')

                # insert informations in treeview ===============================
                self.search_cashManagement(treeview, type_search='last', save_seacrh=False)

                # shoe message of concluded
                self.message_window(1, 'Concluído', messagein=f'Cadastro(s) deletado(s) com sucesso')
        else:
            self.message_window(3, 'Sem seleção', 'Selecione algum item na lista para deletar')

    def pick_informations_for_cash(self, entrys, date=datetime.today().strftime("%m/%Y")):
        # deleting informations of entrys ============================================
        for entry in entrys[0:7]:
            if isinstance(entry, CTkComboBox):
                entry.set('')
            else:
                entry.delete(0, END)
        # informations of schedule and sold ===========================================
        informationsSchedule = self.dataBases['schedule'].searchDatabase(f'SELECT * FROM Agenda WHERE pagamento LIKE "%{date}%"')
        # treating informatios of day for entrys
        informationsOfDay = [
            len([client for client in informationsSchedule if client[3] not in ['NÃO FOI PAGO', 'SEM PAGAMENTO', '']]),
            self.treating_numbers(
                type_treating=4,
                values=[
                    self.treating_numbers(type_treating=2, values=[value for value in informationsSchedule if value[3] in ['CARTÃO', 'CARTAO']], ide=2),
                ]
            ),
            self.treating_numbers(
                type_treating=4,
                values=[
                    self.treating_numbers(type_treating=2, values=[value for value in informationsSchedule if value[3] in ['DINHEIRO']], ide=2),
                ]
            ),
            self.treating_numbers(
                type_treating=4,
                values=[
                    self.treating_numbers(type_treating=2, values=[value for value in informationsSchedule if value[3] in ['TRANSFERÊNCIA', 'TRANSFERENCIA']], ide=2),
                ]
            ),
            self.treating_numbers(
                type_treating=4,
                values=[
                    self.treating_numbers(type_treating=2, values=[value for value in informationsSchedule if value[3] in ['NOTA', 'FIADO', 'NOTINHA']], ide=2),
                ]
            ),
        ]
        # inserting informations of day ====================================
        for index, information in enumerate(informationsOfDay):
            entrys[index].insert(0, information)
        entrys[5].insert(0, self.treating_numbers(type_treating=4, values=informationsOfDay[1:]))
        entrys[6].insert(0, date)

    def create_pdf_cashManagement(self, treeview):
        if informationsTreeview := self.pick_informations_treeview(treeview):
            # informations ============================================
            for information in informationsTreeview:
                tableWithInformationsCashManagementTreeview1.append(information[0:7])
            for information in informationsTreeview:
                tableWithInformationsCashManagementTreeview2.append(information[7:13])
            tableWithInformationsComplementaryCashManagement.append(self.informations_supplementarys_cashManagement(informationsTreeview))

            # create tables ===========================================
            table1 = Table(tableWithInformationsCashManagementTreeview1)
            table1.setStyle(TableStyle(styleTableInformationsTreeview))
            table2 = Table(tableWithInformationsCashManagementTreeview2)
            table2.setStyle(TableStyle(styleTableInformationsTreeview))
            table3 = Table(tableWithInformationsComplementaryCashManagement)
            table3.setStyle(TableStyle(styleTableInformationsComplementary))

            # reseting tables =========================================
            del tableWithInformationsCashManagementTreeview1[2:]
            del tableWithInformationsCashManagementTreeview2[2:]
            del tableWithInformationsComplementaryCashManagement[2:]

            # creating pdf ============================================
            saveDirectory = self.create_pdf(treeview, [table1, table2, table3])
            if saveDirectory is not None:
                self.message_window(1, 'Concluído', messagein=f'O arquivo foi salvo em "{saveDirectory}"')
        else:
            self.message_window(2, 'Sem registro', 'A tabela esta vazia')

    def message_informations_cashManagement(self, treeview):
        if informationsTreeview := self.pick_informations_treeview(treeview):
            # show menssege ================================================
            self.message_window(
                1,
                'Informações sobre a tabela',
                messageCashManagement.format(*self.informations_supplementarys_cashManagement(informationsTreeview))
            )
        else:
            self.message_window(2, 'Sem registro', 'A tabela esta vazia')

    def informations_supplementarys_cashManagement(self, informations):
        # píck up informations supplementarys ==============================
        totalClients = self.treating_numbers(values=[client[1] for client in informations], type_treating=6)
        totalReceived = self.treating_numbers(values=[client[10] for client in informations], type_treating=4)
        totalExit = []
        for information in informations:
            totalExit.append(self.treating_numbers(values=[information[6], information[7], information[8], information[9]], type_treating=4))
        totalExit = self.treating_numbers(values=totalExit, type_treating=4)
        return [totalClients, totalReceived, totalExit]


class FunctionsOfLogin(GeneralFunctions):

    def register_users(self, treeview, informations):
        if self.validation(informations, 5):
            if self.message_window(4, 'Comfimação', f'Finalisar o cadastro de {informations[0].title()}?'):
                # informations of treeview ====================
                self.dataBases['informations'].crud(
                    registerUsers.format(
                        informations[0], self.criptography.crypt(informations[1], type_cryptography="hash"),
                        'NORMAL' if informations[2] == '' else informations[2]
                    )
                )
                # deleting and inserting informations in treeview ===============================
                self.search_users(treeview, informations, 'all', save_seacrh=False)
        else:
            # show message error =====================================================
            self.message_window(3, 'Erro', 'Verifique se os campos estão preenchidos ou corretos')

    def search_users(self, treeview=None, informations=None, type_search='new', save_seacrh=True, insert=True):
        # save last search ============================================
        if save_seacrh:
            self.lastSearch['users'] = searchUsers.format(
                informations[0].upper(), informations[2].upper()
            )

        # pick up informations =========================================
        informationsDatabase = []
        match type_search:
            case 'new':
                informationsDatabase = self.dataBases['informations'].searchDatabase(
                    searchUsers.format(
                        informations[0].upper(), informations[2].upper()
                    )
                )
            case 'last':
                informationsDatabase = self.dataBases['informations'].searchDatabase(self.lastSearch['users'])
            case 'all':
                informationsDatabase = self.dataBases['informations'].searchDatabase(searchAll.format('Usuários'))

        if insert:
            # deleting and inserting informations in treeview ===============================
            self.delete_informations_treeview(treeview, 'users')
            self.insert_treeview_informations(treeview, informationsDatabase, 'users')
        else:
            return informationsDatabase

    def update_users(self, treeview, informations):
        if self.validation([informations[0], informations[2]], 5):
            # update informations =========================================
            if treeview.selection():
                informationsDataBase = self.dataBases['informations'].crud(
                    updateUsers.format(
                        informations[0], self.criptography.crypt(informations[1], type_cryptography="hash"), informations[2].upper(),
                        self.selection_treeview(treeview)[0][0]
                    )
                )
                # delete informations of treeview ==============================
                self.delete_informations_treeview(treeview, 'users')

                # insert informations in treeview ===============================
                self.search_users(treeview, informations, 'last', save_seacrh=False)

                # show message of concluded
                self.message_window(1, 'Concluído', messagein=f'Usuário {informations[0].title()} atualizado com sucesso')
            else:
                self.message_window(3, 'Sem seleção', 'Selecione algum item na lista para atualizar')
        else:
            # show message error =====================================================
            self.message_window(3, 'Erro', 'Verifique se os campos estão preenchidos ou corretos')

    def delete_users(self, treeview):
        if treeview.selection():
            # deleting inforations =======================================
            delete = self.delete_information(treeview, 'informations', 'Usuários')

            if delete:
                # delete informations of treeview ==============================
                self.delete_informations_treeview(treeview, 'users')

                # insert informations in treeview ===============================
                self.search_users(treeview, type_search='last', save_seacrh=False)

                # shoe message of concluded
                self.message_window(1, 'Concluído', messagein=f'Pagamentos(s) deletado(s) com sucesso')
        else:
            self.message_window(3, 'Sem seleção', 'Selecione algum item na lista para deletar')

    def password_window(self, function, parameter):
        self.passwordWindow = Toplevel()
        self.passwordWindow.title('Senha de administrador - BlackBelt-Pro')
        width = self.passwordWindow.winfo_screenwidth()
        height = self.passwordWindow.winfo_screenheight()
        posx = width / 2 - 340 / 2
        posy = height / 2 - 270 / 2
        self.passwordWindow.geometry('340x270+%d+%d' % (posx, posy))
        self.passwordWindow.maxsize(340, 270)
        self.passwordWindow.config(bg='#FFFFFF')
        self.passwordWindow.focus_force()
        self.passwordWindow.iconphoto(False, PhotoImage(file=self.dataBases['config'].searchDatabase('SELECT * FROM Logo')[0][0]))

        # frame of inputs ===========================================================
        frameInputs = self.frame(self.passwordWindow, 0, 0.01, 1, 0.985)

        # photo ---------------------------------------
        self.labelPadlock = self.labels(self.passwordWindow, '', 0.3, 0.1, width=0.4, height=0.2, photo=self.image('assets/icon_password.png', (56, 56))[0], position=CENTER)

        # password ------------------------------------
        password = self.entry(frameInputs, 0.1, 0.35, 0.8, 0.16, type_entry='entry', show='*')
        password.focus_force()

        # visibility ---------------------------------
        visibilityPasswordBtn = self.button(
            self.passwordWindow, '', 0.76, 0.38, 0.13, 0.1, photo=self.image('assets/icon_eyeClose.png', (26, 26))[0],
            type_btn='buttonPhoto', hover_cursor='white', function=lambda: self.toggle_visibility(password, visibilityPasswordBtn)
        )
        password.bind('<KeyPress>', lambda e: [password.configure(fg_color='#FFFFFF'), visibilityPasswordBtn.configure(fg_color='#FFFFFF')])
        password.bind('<Return>', lambda e: self.validating_user([password, visibilityPasswordBtn], function, 'password', parameters=parameter, window_password=self.passwordWindow))

        # comfirm -------------
        comfirmBtn = self.button(
            self.passwordWindow, 'Confirmar', 0.1, 0.56, 0.8, 0.17,
            function=lambda: self.validating_user([password, visibilityPasswordBtn], function, 'password', parameters=parameter, window_password=self.passwordWindow)
        )

        # cancel -------------
        cancelBtn = self.button(self.passwordWindow, 'Cancelar', 0.1, 0.76, 0.8, 0.17, function=lambda: self.passwordWindow.destroy())

    def validating_user(self, entrys, function, type_password, parameters, window_password=None):
        # searching infomation of user =========================
        informations = []
        match type_password:
            case 'login':
                informations = self.dataBases['informations'].searchDatabase(
                    f'SELECT * FROM Usuários WHERE nome = "{entrys[0].get()}" AND senha = "{self.criptography.crypt(entrys[1].get(), type_cryptography="hash")}"'
                )
            case 'password':
                informations = self.dataBases['informations'].searchDatabase(
                    f'SELECT * FROM Usuários WHERE senha = "{self.criptography.crypt(entrys[0].get(), type_cryptography="hash")}" AND nivel = "ADMINISTRADOR"'
                )

        if informations:
            if type_password == 'password':
                window_password.destroy()
            return function(*parameters.values())
        else:
            colors = ['#f07f7f', '#FFFFFF']
            if len(entrys) > 2:
                entrys[0].configure(fg_color='#f07f7f')
                entrys[1].configure(fg_color='#f07f7f')
                entrys[2].configure(fg_color='#f07f7f')
            else:
                entrys[0].configure(fg_color='#f07f7f')
                entrys[1].configure(fg_color='#f07f7f')

    def toggle_visibility(self, entry, button):
        current_show = entry.cget("show")
        if current_show == "*":
            entry.configure(show="")
            button.configure(image=self.image('assets/icon_eyeOpen.png', (26, 26))[0])
        else:
            entry.configure(show="*")
            button.configure(image=self.image('assets/icon_eyeClose.png', (26, 26))[0])

    def open_software(self, event):
        self.loginWindow.withdraw()
        self.main_window()


class FunctionsOfConfigurations(GeneralFunctions):

    def colorPicker(self, widget, type_color, entry_color, color_picker='yes'):
        if self.openColorPicker or color_picker == 'no':
            selectorColor = ''
            match color_picker:
                case 'yes':
                    # pick color ==========================
                    self.openColorPicker = False
                    selectorColor = askcolor(initialcolor=entry_color.get())[1]
                    self.openColorPicker = True
                case 'no':
                    selectorColor = entry_color.get()
            if selectorColor is not None:
                # update widget and color of entry =======================
                match type_color:
                    case 'fg_color':
                        try:
                            widget.configure(fg_color=selectorColor)
                        except TclError:
                            self.message_window(3, 'Erro', 'Valor hexadecimal inválido ou cor não existente')
                    case 'text_color':
                        try:
                            widget.configure(text_color=selectorColor)
                        except TclError:
                            self.message_window(3, 'Erro', 'Valor hexadecimal inválido ou cor não existente')
                    case 'border_color':
                        try:
                            widget.configure(border_color=selectorColor)
                        except TclError:
                            self.message_window(3, 'Erro', 'Valor hexadecimal inválido ou cor não existente')
                    case 'hover_color':
                        try:
                            widget.configure(hover_color=selectorColor)
                        except TclError:
                            self.message_window(3, 'Erro', 'Valor hexadecimal inválido ou cor não existente')
                    case 'tag1':
                        try:
                            widget.tag_configure('oddrow', background=selectorColor)
                        except TclError:
                            self.message_window(3, 'Erro', 'Valor hexadecimal inválido ou cor não existente')
                    case 'tag2':
                        try:
                            widget.tag_configure('evenrow', background=selectorColor)
                        except TclError:
                            self.message_window(3, 'Erro', 'Valor hexadecimal inválido ou cor não existente')
                    case 'text_color_treeview':
                        try:
                            self.style_treeview.configure('Treeview', foreground=selectorColor)
                        except TclError:
                            self.message_window(3, 'Erro', 'Valor hexadecimal inválido ou cor não existente')
                entry_color.delete(0, END)
                entry_color.insert(0, selectorColor)
        else:
            # show message error =====================================================
            self.message_window(2, 'Já em uso', 'O seletor de cores ainda está em uso')

    def save_configs(self):
        # saving configs of buttons ================================
        framesOfWidgets = [self.frameForButtons, self.frameForFrames, self.frameForTabview, self.frameForTreeview, self.frameForEntrys, self.frameForLabels]
        colors = []
        for index, frame in enumerate(framesOfWidgets):
            for widget in frame.winfo_children():
                if isinstance(widget, CTkEntry):
                    colors.append(widget.get())
            match index:
                case 0:
                    self.dataBases['config'].crud(f'UPDATE Botões SET cor_de_fundo="{colors[0]}", cor_de_texto="{colors[1]}", cor_da_borda="{colors[2]}", cor_do_hover="{colors[3]}"')
                case 1:
                    self.dataBases['config'].crud(f'UPDATE Frames SET cor_de_fundo="{colors[0]}", cor_da_borda="{colors[1]}"')
                case 2:
                    self.dataBases['config'].crud(f'UPDATE Tabviews SET cor_de_fundo="{colors[0]}", cor_da_borda="{colors[1]}"')
                case 3:
                    self.dataBases['config'].crud(f'UPDATE Treeviews SET cor_da_linha1="{colors[0]}", cor_da_linha2="{colors[1]}", cor_de_texto="{colors[2]}"')
                case 4:
                    self.dataBases['config'].crud(f'UPDATE Entrys SET cor_de_fundo="{colors[0]}", cor_de_texto="{colors[1]}", cor_da_borda="{colors[2]}"')
                case 5:
                    self.dataBases['config'].crud(f'UPDATE Labels SET cor_de_texto1="{colors[0]}", cor_de_texto2="{colors[1]}"')
            del colors[0:]
        self.message_window(1, 'Concluído', 'Reabra o programa para efetuar as alterações')

    def load_configs(self):
        framesOfWidgets = [self.frameForButtons, self.frameForFrames, self.frameForTabview, self.frameForTreeview, self.frameForEntrys, self.frameForLabels]
        entrys = []
        colors = []
        for indexf, frame in enumerate(framesOfWidgets):
            for widget in frame.winfo_children():
                if isinstance(widget, CTkEntry):
                    entrys.append(widget)
            match indexf:
                case 0:
                    colors = self.dataBases['config'].searchDatabase('SELECT * FROM Botões')[0]
                case 1:
                    colors = self.dataBases['config'].searchDatabase('SELECT * FROM Frames')[0]
                case 2:
                    colors = self.dataBases['config'].searchDatabase('SELECT * FROM Tabviews')[0]
                case 3:
                    colors = self.dataBases['config'].searchDatabase('SELECT * FROM Treeviews')[0]
                case 4:
                    colors = self.dataBases['config'].searchDatabase('SELECT * FROM Entrys')[0]
                case 5:
                    colors = self.dataBases['config'].searchDatabase('SELECT * FROM Labels')[0]
            for indexc, color in enumerate(colors):
                entrys[indexc].insert(0, color)
            del entrys[0:]
            colors = []
